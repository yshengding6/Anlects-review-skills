# -*- coding: utf-8 -*-
"""
古文典籍文献综述提示词格式转换脚本
将 Markdown 文件转换为 PDF 和 DOCX 格式
"""

import os
import sys
from pathlib import Path

# 设置输出编码为UTF-8
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_docx(md_file_path, output_path):
    """将 Markdown 文件转换为 DOCX 格式"""
    print(f"正在生成 DOCX 文件: {output_path}")

    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建 Word 文档
    doc = Document()

    # 设置文档默认字体（中文字体）
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 逐行处理 Markdown 内容
    lines = md_content.split('\n')

    for line in lines:
        line = line.rstrip()

        # 空行
        if not line:
            doc.add_paragraph()
            continue

        # 标题处理
        if line.startswith('#'):
            level = line.count('#')
            title = line.lstrip('#').strip()
            heading = doc.add_heading(title, level=min(level, 3))

            # 设置标题字体
            for run in heading.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 引用块
        elif line.startswith('>'):
            text = line.lstrip('>').strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.italic = True

        # 分割线
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 50).font.color.rgb = RGBColor(128, 128, 128)

        # 列表项
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')

        # 数字列表
        elif line.startswith(tuple([f'{i}. ' for i in range(1, 10)])):
            p = doc.add_paragraph(line, style='List Number')

        # 表格（简化处理）
        elif '|' in line and line.count('|') >= 2:
            # 这里只做简单处理，完整表格需要更复杂的解析
            p = doc.add_paragraph(line)

        # 加粗文本
        elif '**' in line:
            parts = line.split('**')
            p = doc.add_paragraph()
            i = 0
            while i < len(parts):
                if i % 2 == 0:
                    p.add_run(parts[i])
                else:
                    run = p.add_run(parts[i])
                    run.bold = True
                i += 1

        # 普通段落
        else:
            doc.add_paragraph(line)

    # 保存 DOCX 文件
    doc.save(output_path)
    print(f"✓ DOCX 文件已生成: {output_path}")


def create_pdf(md_file_path, output_path):
    """将 Markdown 文件转换为 PDF 格式"""
    print(f"正在生成 PDF 文件: {output_path}")

    try:
        # 尝试使用 weasyprint
        from weasyprint import HTML, CSS
        from markdown import markdown

        # 读取 Markdown 文件
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 转换为 HTML
        html_content = markdown(md_content, extensions=['tables', 'fenced_code'])

        # 添加 CSS 样式
        css_style = """
        @import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;700&display=swap');

        body {
            font-family: 'Noto Serif SC', 'SimSun', '宋体', serif;
            font-size: 12pt;
            line-height: 1.6;
            margin: 2cm;
            color: #333;
        }

        h1 {
            font-size: 24pt;
            color: #1a1a1a;
            border-bottom: 2px solid #333;
            padding-bottom: 0.5em;
            margin-top: 1em;
            margin-bottom: 1em;
        }

        h2 {
            font-size: 18pt;
            color: #2a2a2a;
            border-bottom: 1px solid #ccc;
            padding-bottom: 0.3em;
            margin-top: 1.5em;
            margin-bottom: 1em;
        }

        h3 {
            font-size: 14pt;
            color: #3a3a3a;
            margin-top: 1.2em;
            margin-bottom: 0.8em;
        }

        h4 {
            font-size: 12pt;
            color: #4a4a4a;
            margin-top: 1em;
            margin-bottom: 0.6em;
        }

        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            font-size: 10pt;
        }

        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }

        th {
            background-color: #f5f5f5;
            font-weight: bold;
        }

        blockquote {
            border-left: 4px solid #666;
            margin: 1em 0;
            padding-left: 1em;
            color: #555;
            font-style: italic;
        }

        code {
            font-family: 'Consolas', 'Courier New', monospace;
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-size: 0.9em;
        }

        pre {
            background-color: #f4f4f4;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
        }

        pre code {
            background-color: transparent;
            padding: 0;
        }

        ul, ol {
            margin: 0.5em 0;
        }

        li {
            margin: 0.3em 0;
        }

        hr {
            border: none;
            border-top: 1px solid #ccc;
            margin: 2em 0;
        }

        strong {
            font-weight: bold;
        }

        em {
            font-style: italic;
        }
        """

        # 创建完整的 HTML 文档
        full_html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>古文典籍文献综述 · 专业版提示词</title>
            <style>{css_style}</style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        # 生成 PDF
        HTML(string=full_html).write_pdf(output_path)
        print(f"✓ PDF 文件已生成: {output_path}")

    except ImportError:
        print("⚠ 缺少 PDF 生成依赖库，请安装: pip install weasyprint markdown")


def main():
    """主函数"""
    # 获取当前脚本所在目录
    script_dir = Path(__file__).parent
    md_file = script_dir / "古文典籍文献综述_标准提示词_v2.0.md"

    if not md_file.exists():
        print(f"错误: 找不到文件 {md_file}")
        return

    # 输出文件路径
    docx_output = script_dir / "古文典籍文献综述_标准提示词_v2.0.docx"
    pdf_output = script_dir / "古文典籍文献综述_标准提示词_v2.0.pdf"

    # 转换为 DOCX
    try:
        create_docx(md_file, docx_output)
    except ImportError:
        print("⚠ 缺少 DOCX 生成依赖库，请安装: pip install python-docx")

    # 转换为 PDF
    try:
        create_pdf(md_file, pdf_output)
    except Exception as e:
        print(f"⚠ PDF 生成失败: {e}")
        print("提示: 如需生成 PDF，请安装: pip install weasyprint markdown")


if __name__ == '__main__':
    main()
